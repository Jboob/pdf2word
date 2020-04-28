import sys
import importlib
from docx import Document
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBoxHorizontal
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed, PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfparser import PDFParser, PDFDocument
importlib.reload(sys)


class PDF2Word:
    def __init__(self, pdf_path):
        # 以二进制读模式打开 PDF
        fp = open(pdf_path, 'rb')
        # 用文件对象来创建一个 PDF 文档分析器
        parser = PDFParser(fp)
        # 创建一个 PDF 文档
        self.doc = PDFDocument()
        # 连接分析器 与文档对象
        parser.set_document(self.doc)
        self.doc.set_parser(parser)

        # 提供初始化密码
        # 如果没有密码 就创建一个空的字符串
        self.doc.initialize()

    # PDF 转 Word
    def pdf_to_word(self, sve_path):
        # 检测文档是否提供 txt 转换，不提供就忽略
        if not self.doc.is_extractable:
            raise PDFTextExtractionNotAllowed
        else:
            # 创建 PDf 资源管理器 来管理共享资源
            rsrcmgr = PDFResourceManager()
            # 创建一个 PDF 设备对象
            laparams = LAParams()
            device = PDFPageAggregator(rsrcmgr, laparams=laparams)
            # 创建一个 PDF 解释器对象
            interpreter = PDFPageInterpreter(rsrcmgr, device)

            # 用来计数页面，图片，曲线，figure，水平文本框等对象的数量
            num_page, num_image, num_curve, num_figure, num_TextBoxHorizontal = 0, 0, 0, 0, 0

            # 首先创建一个文档对象
            document = Document()
            # 循环遍历列表，每次处理一个page的内容
            for page in self.doc.get_pages():  # doc.get_pages() 获取page列表
                num_page += 1  # 页面增一
                interpreter.process_page(page)
                # 接受该页面的LTPage对象
                layout = device.get_result()
                for x in layout:
                    if isinstance(x, LTTextBoxHorizontal):  # 获取文本内容
                        results = x.get_text()
                        document.add_paragraph(results)
            document.save(sve_path)


if __name__ == '__main__':
    pdf_path = '阿里巴巴Java开发手册1.4.0.pdf'
    covertFile = PDF2Word(pdf_path)
    covertFile.pdf_to_word('阿里巴巴Java开发手册1.4.0.docx')
